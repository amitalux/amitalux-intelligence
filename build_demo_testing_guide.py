from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUTPUT = "Amitalux Intelligence_Demo_Testing_Guide.docx"
DEMO_URL = "https://6a448e25906e9c02f96fdb0e--amitalux-customer-care.netlify.app"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(9.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run_element = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "4D78C7")
    properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(underline)
    run_element.append(properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run_element.append(text_element)
    hyperlink.append(run_element)
    paragraph._p.append(hyperlink)


def add_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.add_run(text)


def add_numbered(document, text):
    paragraph = document.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.add_run(text)


def add_section(document, title):
    paragraph = document.add_paragraph()
    paragraph.style = document.styles["Heading 1"]
    paragraph.add_run(title)
    return paragraph


def add_subsection(document, title):
    paragraph = document.add_paragraph()
    paragraph.style = document.styles["Heading 2"]
    paragraph.add_run(title)
    return paragraph


def add_callout(document, title, body):
    table = document.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.3)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "E7F0FF")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(4)
    title_run = paragraph.add_run(title)
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run.font.size = Pt(10.5)
    body_paragraph = cell.add_paragraph()
    body_paragraph.paragraph_format.space_after = Pt(0)
    body_run = body_paragraph.add_run(body)
    body_run.font.name = "Arial"
    body_run.font.size = Pt(10)
    document.add_paragraph()


def build_document():
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color, before, after in [
        ("Heading 1", 16, "20232D", 14, 6),
        ("Heading 2", 12.5, "4D78C7", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run("Amitalux Intelligence Demo Testing Guide")
    run.font.name = "Arial"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(32, 35, 45)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    subtitle_run = subtitle.add_run("A simple walkthrough for reviewing the live SaaS prototype")
    subtitle_run.font.name = "Arial"
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = RGBColor(104, 114, 132)

    link_paragraph = document.add_paragraph()
    link_paragraph.add_run("Demo link: ").bold = True
    add_hyperlink(link_paragraph, DEMO_URL, DEMO_URL)

    add_callout(
        document,
        "Purpose of this test",
        "This guide helps a reviewer experience Amitalux Intelligence as a customer, an agent, and an admin. The goal is to confirm that the product feels like functioning customer service software, not a static website.",
    )

    add_section(document, "Before You Start")
    add_bullet(document, "Open the demo link in Chrome, Safari, or Edge.")
    add_bullet(document, "If something looks stale, refresh the page once.")
    add_bullet(document, "Use the left sidebar to move between Customer Portal, Agent Portal, Admin Portal, Reports, and Profiles.")
    add_bullet(document, "The prototype uses demo customer records and demo business data. Do not enter private real customer information yet.")

    add_section(document, "Recommended 10 Minute Test Path")
    steps = [
        "Start on YC Demo and read the product positioning.",
        "Open Customer Portal and send a customer message.",
        "Open Agent Portal and confirm the new message appears in the queue.",
        "Run AI analysis, approve the reply, send it, and close the case.",
        "Open Admin Portal and review operating health, queue, stale data, history, and agent performance.",
        "Open Reports and check volume, queue, CSAT, agent report, and case progression.",
        "Open Customer Profile and Agent Profile to confirm the support context is available.",
    ]
    for step in steps:
        add_numbered(document, step)

    add_section(document, "Test Scenarios")
    table = document.add_table(rows=1, cols=4)
    table.autofit = False
    widths = [1.25, 1.65, 1.7, 1.7]
    headers = ["Area", "Action", "Expected Result", "What To Notice"]
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.width = Inches(widths[index])
        set_cell_shading(cell, "F2F4F7")
        set_cell_text(cell, header, bold=True)

    rows = [
        (
            "Customer Portal",
            "Send the default Maya message.",
            "The send status changes and the message appears in the customer thread.",
            "The customer-facing experience should feel simple and calm.",
        ),
        (
            "Agent Portal",
            "Select the newest case and run AI analysis.",
            "The case status changes and a personal reply draft appears.",
            "The AI should sound human and should use customer context.",
        ),
        (
            "Agent Portal",
            "Approve, send, then close the case.",
            "The timeline updates and the case moves through the workflow.",
            "This should feel like actual case handling.",
        ),
        (
            "Admin Portal",
            "Review health, queue, stale data, and history.",
            "Counts, red/yellow/green signals, and history should reflect work done.",
            "Admins should quickly see what needs attention.",
        ),
        (
            "Reports",
            "Review volume, queue, CSAT, agent report, and progression.",
            "The dashboard should answer operational questions clearly.",
            "This is where a buyer checks if Amitalux Intelligence can run the business day.",
        ),
        (
            "Profiles",
            "Open Customer Profile and Agent Profile.",
            "Customer context, memory, risk, permissions, and coaching are visible.",
            "A support agent should not need to hunt through separate systems.",
        ),
    ]

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].width = Inches(widths[index])
            set_cell_text(cells[index], value)

    add_section(document, "Customer Portal Test")
    add_numbered(document, "Click Customer Portal.")
    add_numbered(document, "Leave the email as maya@example.com or enter sophia@example.com or jordan@example.com.")
    add_numbered(document, "Choose Chat or Email.")
    add_numbered(document, "Send a message such as: I still need help with my late fee. Can someone please look at this today?")
    add_numbered(document, "Confirm that the conversation thread and customer-facing email section update.")

    add_subsection(document, "Expected result")
    add_bullet(document, "The customer sees a simple thread instead of a complex ticket system.")
    add_bullet(document, "The message enters the operational queue for the business.")
    add_bullet(document, "If online sync pauses, the app should show a short status message and keep the demo usable.")

    add_section(document, "Agent Portal Test")
    add_numbered(document, "Click Agent Portal.")
    add_numbered(document, "Select the newest or most urgent case from Open conversations.")
    add_numbered(document, "Click Run AI analysis.")
    add_numbered(document, "Read the AI-generated reply and the case context.")
    add_numbered(document, "Click Approve reply, Send reply, and Close case.")

    add_subsection(document, "Expected result")
    add_bullet(document, "The reply should acknowledge the customer personally.")
    add_bullet(document, "The agent should see channel, sentiment, context, timeline, and customer details.")
    add_bullet(document, "The case should progress through Received, Analyzed, Approved, Sent, and Closed.")

    add_section(document, "Admin Portal Test")
    add_numbered(document, "Click Admin Portal.")
    add_numbered(document, "Review the operating health card.")
    add_numbered(document, "Check Open conversations, Agent performance, Stale data alerts, and Recent history.")
    add_numbered(document, "Look for red, yellow, and green status signals.")

    add_subsection(document, "Expected result")
    add_bullet(document, "An admin should quickly understand volume, queue risk, customer risk, and team performance.")
    add_bullet(document, "Stale customer data should be visible without needing to search for it.")
    add_bullet(document, "Recent history should show what happened during the demo.")

    add_section(document, "Reports Test")
    add_bullet(document, "Open Reports and review total volume, volume change, queue total, queue risk, CSAT score, CSAT change, agent report, and case progression.")
    add_bullet(document, "Confirm that the page answers the questions: How is volume changing? How is the queue? How are agents doing? How is CSAT? How is the case progressing?")

    add_section(document, "Profiles Test")
    add_bullet(document, "Open Customer Profile to review customer memory, Square-style context, CRM-style context, risk, prior resolutions, and available solutions.")
    add_bullet(document, "Open Agent Profile to review role, load, permissions, tone, coaching, handled count, CSAT, and average reply time.")

    add_section(document, "Feedback Checklist")
    checklist = [
        "Did the product feel like real customer service software?",
        "Could a small business owner understand the value quickly?",
        "Did the AI reply sound personal and human?",
        "Did the agent have enough information to help without asking the customer to repeat everything?",
        "Did the admin dashboard answer operational questions clearly?",
        "Was anything confusing, too static, or missing?",
        "What would make this feel sellable as a SaaS product?",
    ]
    for item in checklist:
        add_bullet(document, item)

    add_section(document, "Known Prototype Boundaries")
    add_bullet(document, "This is a working prototype, not the final paid production architecture.")
    add_bullet(document, "Real login, tenant billing, production database, real email sending, real SMS, and live Square or Salesforce connections are next-stage build items.")
    add_bullet(document, "The current demo is meant to prove the product workflow, customer experience, agent workflow, admin visibility, and SaaS concept.")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Amitalux Intelligence | Demo Testing Guide")
    footer_run.font.name = "Arial"
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(104, 114, 132)

    document.save(OUTPUT)


if __name__ == "__main__":
    build_document()
